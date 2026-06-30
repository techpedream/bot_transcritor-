e v     import streamlit as st
from openai import OpenAI
from docx import Document
import tempfile
import os
import subprocess
import json
from datetime import date
from pathlib import Path
import shutil

# --- Configuração da API ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    st.error("⚠️ API Key da OpenAI não encontrada.")
    st.stop()

# Define a variável de ambiente para garantir compatibilidade com o client
os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY

# Inicializa o client da OpenAI
client = OpenAI()

# --- Parametrização de áudio ---
SEGUNDO_POR_PARTE = 600  # 10 minutos
TARGET_SR = 16000        # 16 kHz
BITRATE = "48k"          # MP3 leve

ALLOWED_TYPES = ["mp3", "m4a", "wav", "ogg"]

def converter_para_mp3_mono16k(caminho_entrada: str) -> str:
    """Converte qualquer áudio para MP3, 16 kHz, mono, bitrate baixo."""
    caminho_saida = Path(caminho_entrada).with_suffix("").as_posix() + "_proc.mp3"
    result = subprocess.run(
        ["ffmpeg", "-y", "-i", caminho_entrada,
         "-ar", str(TARGET_SR), "-ac", "1", "-b:a", BITRATE, caminho_saida],
        capture_output=True
    )
    if result.returncode != 0:
        stderr = result.stderr.decode("utf-8", errors="replace")
        raise RuntimeError(f"ffmpeg erro (código {result.returncode}):\n{stderr[-1000:]}")
    return caminho_saida

def duracao_segundos(caminho: str) -> float:
    """Retorna a duração do áudio em segundos via ffprobe."""
    try:
        result = subprocess.run(
            ["ffprobe", "-v", "quiet", "-print_format", "json", "-show_format", caminho],
            capture_output=True, text=True, check=True
        )
        info = json.loads(result.stdout)
        return float(info["format"]["duration"])
    except Exception:
        return 0.0

def fatiar_audio(path_mp3: str, segundos_por_parte: int = SEGUNDO_POR_PARTE) -> list:
    """Fatia o MP3 em partes de até `segundos_por_parte`. Retorna lista ordenada com caminhos."""
    try:
        dur = duracao_segundos(path_mp3)
        if dur == 0:
            return [path_mp3]

        partes = []
        out_dir = Path(tempfile.mkdtemp(prefix="chunks_"))
        start = 0
        i = 0
        while start < dur:
            i += 1
            caminho_parte = str(out_dir / f"parte_{i:03d}.mp3")
            subprocess.run(
                ["ffmpeg", "-y", "-i", path_mp3,
                 "-ss", str(start), "-t", str(segundos_por_parte),
                 "-b:a", BITRATE, caminho_parte],
                check=True, capture_output=True
            )
            partes.append(caminho_parte)
            start += segundos_por_parte

        return partes if partes else [path_mp3]
    except Exception as e:
        print(f"Erro ao fatiar áudio: {e}")
        return [path_mp3]


# -------- Transcrição (Whisper) --------
def transcrever_um_arquivo(caminho_audio: str) -> str:
    """Transcreve um arquivo individual com Whisper-1 e retorna texto."""
    with open(caminho_audio, "rb") as f:
        tr = client.audio.transcriptions.create(
            model="whisper-1",
            file=f
        )
    return tr.text or ""

def transcrever_em_partes(partes: list) -> str:
    """Transcreve cada parte e concatena com separadores claros."""
    textos = []
    for idx, p in enumerate(partes, start=1):
        try:
            texto = transcrever_um_arquivo(p).strip()
            if texto:
                textos.append(f"[PARTE {idx}]\n{texto}")
        except Exception as e:
            st.warning(f"Falha ao transcrever a parte {idx}: {e}")
    return "\n\n".join(textos).strip()

# -------- App --------
st.title("🧠 Gerador de Relatório de Psicoterapia")

audio_file = st.file_uploader("🎤 Envie o áudio da sessão", type=None)
nome_paciente = st.text_input("🧍 Nome do Paciente")
idade_paciente = st.text_input("📅 Idade do Paciente")
numero_sessao = st.text_input("🔢 Número da Sessão")
nome_psicologo = st.text_input("👩‍⚕️ Acadêmico(a)")
nome_supervisor = st.text_input("🧑‍🏫 Supervisor(a)")
data_sessao = st.date_input("📅 Data da Sessão", value=date.today())

if st.button("🚀 Gerar Relatório") and audio_file:
    log = st.container()

    def log_ok(msg):
        log.success(msg)

    def log_err(msg):
        log.error(msg)
        st.stop()

    # 1) Salvar upload
    with st.spinner("Salvando arquivo..."):
        ext = os.path.splitext(audio_file.name)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(audio_file.read())
            original_path = tmp.name
        tamanho_mb = os.path.getsize(original_path) / 1024 / 1024
        log_ok(f"[1/5] Arquivo salvo: {audio_file.name} ({tamanho_mb:.1f} MB)")

    # 2) Converter p/ MP3 leve (16kHz, mono)
    with st.spinner("Convertendo áudio (16 kHz, mono, MP3)..."):
        try:
            mp3_path = converter_para_mp3_mono16k(original_path)
        except Exception as e:
            log_err(f"[2/5] Erro na conversão ffmpeg: {e}")
        if not mp3_path or not os.path.exists(mp3_path):
            log_err("[2/5] Conversão falhou — verifique se o ffmpeg está instalado no servidor.")
        tamanho_mp3 = os.path.getsize(mp3_path) / 1024 / 1024
        log_ok(f"[2/5] Áudio convertido: {tamanho_mp3:.1f} MB")

    # 3) Fatiar se for longo
    with st.spinner("Verificando duração e fatiando, se necessário..."):
        dur = duracao_segundos(mp3_path)
        minutos = int(dur // 60)
        if dur > SEGUNDO_POR_PARTE:
            partes = fatiar_audio(mp3_path, SEGUNDO_POR_PARTE)
            log_ok(f"[3/5] Duração: {minutos} min — fatiado em {len(partes)} parte(s)")
        else:
            partes = [mp3_path]
            log_ok(f"[3/5] Duração: {minutos} min — sem necessidade de fatiar")

    # 4) Transcrever tudo
    with st.spinner(f"Transcrevendo {len(partes)} parte(s) com Whisper... ⏳"):
        try:
            texto_transcrito = transcrever_em_partes(partes)
            if not texto_transcrito:
                log_err("[4/5] Transcrição ficou vazia — o Whisper não retornou texto.")
            log_ok(f"[4/5] Transcrição concluída: {len(texto_transcrito)} caracteres")
        except Exception as e:
            log_err(f"[4/5] Erro na transcrição Whisper: {e}")

    # 5) Gerar relatório com GPT
    with st.spinner("Gerando relatório com IA... ✨"):
        prompt = f"""
Você é um psicólogo clínico com foco em descrição comportamental e análise funcional.

Sua tarefa é gerar um relatório descritivo e operacional a partir da transcrição da sessão, que está dividida em partes [PARTE 1], [PARTE 2], etc.

⚠️ Regras importantes:
- Considere TODAS as partes igualmente, da primeira até a última. Não priorize apenas o início.
- Extraia informações relevantes de CADA parte antes de consolidar.
- Se um tema aparece só no final, registre no relatório também.
- Use apenas informações presentes na transcrição. Não invente dados. Não diagnostique.
- Evite termos circulares e “ficções explicativas”.
- Quando a informação não estiver presente, escreva literalmente “não mencionado”.

Estilo e regras de escrita:
- Priorize comportamentos observáveis, contexto e consequências (ABC).
- Faça o ABC do COMPORTAMENTO RELATADO, ao invés do ABC do RELATO. Importante entender o que o paciente traz no relato. A contingência direta descrita e não a contingencia verbal apenas.
- Estados internos sempre como relato: “o paciente relata…”.
- Frases curtas, voz ativa, linguagem clara.
- Hipóteses devem ser marcadas como “Hipótese de trabalho”, com evidências e nível de confiança (Baixo/Moderado/Alto).
- Não cite siglas de abordagens terapêuticas.

FORMATO DE SAÍDA (obrigatório):

REGISTRO DOCUMENTAL
Paciente: {nome_paciente}
Idade: {idade_paciente}
Data: {data_sessao.strftime('%d/%m/%Y')} Sessão: {numero_sessao}
Acadêmicos: {nome_psicologo}
Supervisor: {nome_supervisor}

1. Relato do paciente (descritivo)

Contextos/Antecedentes (A): [situações, locais, pessoas, eventos “antes de”]

Respostas/Comportamentos (B): [o que fez/disse; frequência/intensidade/duração se citadas]

Consequências (C): [o que ocorreu depois; reações de outros; efeitos imediatos]

Estados internos relatados: [sentimentos/pensamentos como relatos, não como causas]

Observação: use trechos curtos entre aspas quando relevante.

2. Análise descritiva e hipóteses de trabalho

Padrões observáveis: [relações A→B→C recorrentes; gatilhos relatados]

Variáveis contextuais relevantes: [horário, ambiente, pessoas, condições]

Lacunas de informação: [perguntas objetivas a esclarecer]

Hipóteses de trabalho (se houver):

Hipótese: [enunciado operacional, sem rótulos diagnósticos]  
Evidências da transcrição: [itens específicos]  
Nível de confiança: [Baixo/Moderado/Alto]  
Alternativas/Falsificadores: [o que mudaria a hipótese]  

3. Exame do estado mental (com base no que foi dito)

Estado emocional atual (última semana): [descrever / não mencionado]  
Aspectos do paciente na entrevista: [descrever / não mencionado]  
Comunicação com os estagiários: [descrever / não mencionado]  
Sentimentos verbalizados ou demonstrados: [descrever / não mencionado]  
Linguagem: [descrever / não mencionado]  

4. Observações gerais

[observações adicionais objetivas; sem rótulos diagnósticos]

5. Sugestões para o próximo atendimento (operacionais)

Liste recomendações comportamentais claras, sem citar abordagens:

[Ação] + [Contexto] + [Frequência/Duração] + [Critério de sucesso].  
Ex.: “Registrar 1 episódio por dia usando ABC em situações de X, por 7 dias.”  
Ex.: “Praticar 2 minutos de respiração diafragma após [evento gatilho] por 1 semana.”  
Ex.: “Planejar 3 atividades prazerosas específicas para [dias/horários], e executar.”  

6. Resumo final (obrigatório)

Queixa principal ou tema da sessão:  
Intervenções realizadas:  
Evolução do paciente:  


=== TRANSCRIÇÃO COMPLETA ===
{texto_transcrito}
"""
        try:
            resposta = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}]
            )
            texto_relatorio = resposta.choices[0].message.content.strip()
        except Exception as e:
            st.error(f"Erro ao gerar relatório: {e}")
            st.stop()

    # 6) Montar DOCX: relatório + transcrição completa
    with st.spinner("Gerando DOCX... 📄"):
        doc = Document()

        # Relatório formatado (linha a linha preserva quebras)
        for linha in texto_relatorio.splitlines():
            doc.add_paragraph(linha)

        # Página nova + transcrição
        doc.add_page_break()
        doc.add_heading("Transcrição Completa", level=1)
        for linha in texto_transcrito.splitlines():
            doc.add_paragraph(linha)

        doc_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
        doc.save(doc_path)

    with open(doc_path, "rb") as f:
        st.success("✅ Relatório gerado!")
        st.download_button(
            "📥 Baixar Relatório",
            f,
            file_name=f"Relatorio_{nome_paciente}.docx"
        )

    # Limpeza básica (opcional)
    try:
        # remove diretórios temporários usados na segmentação
        for p in Path(tempfile.gettempdir()).glob("chunks_*"):
            if p.is_dir():
                shutil.rmtree(p, ignore_errors=True)
    except Exception:
        pass









